from docx import Document
import sys
import os
import argparse


def get_notes(doc):
    notes = {}
    if doc.endnotes_part is not None:
        for n in doc.endnotes_part.notes:
            notes[n.id] = n

    note_ids = set(notes.keys())
    return notes, note_ids


def get_paragraphs(doc, chapter=None):
    """
    Collects all of the paragraphs from a Document() object that have references in them.

    Returns a list of (ref, paragraph) paris, where paragraph is the paragraph that contains
    the given reference.

    These are returned in the same order as they occur in the text.
    """
    runs = []
    current_chapter = 0
    for p in doc.paragraphs:
        if p.text.startswith("Chapter"):
            current_chapter = int(p.text.lstrip("Chapter").split(":")[0])
        for r in p.runs:
            if len(r.endnote_references) > 0:
                if chapter is None or current_chapter == chapter:
                    runs.append((r, p))
    return runs


def ref_to_run_ix(id, runs):
    """
    Given a collection of runs which may or may not have endnote references, it will find
    the one that has the given endnote reference id.
    """
    for ix, run in enumerate(runs):
        if len(run.endnote_references) > 0 and run.endnote_references[0].id == id:
            return ix
    return -1


def get_previous_words(ix, runs, num_prev):
    """
    Collects the words that lead up to a reference from the main text. Does some very basic
    selection to arrange to choose e.g. the start of a sentence, or quote
    """
    if num_prev <= 0:
        return ""

    pre_text = "".join([r.text for r in runs[: ix + 1]]).split()
    num_words = min(num_prev, len(pre_text))

    out_text = " ".join(pre_text[-num_words:])
    ix = out_text.rfind(". ")
    if -1 != ix:
        out_text = out_text[ix + 2 :]

    ix = out_text.rfind("“")
    if -1 != ix:
        out_text = out_text[ix:]

    out_text = out_text.strip()

    if out_text[0].islower():
        out_text = "..." + out_text

    if out_text[-1] not in ".”),;":
        out_text += "..."
    else:
        out_text = out_text[:-1]

    return out_text


def ref_to_note(notes):
    """
    Converts a reference (i.e. a pair of a ref + its paragraph) to the text of its endnote.

    Used for sorting:

    pars = sorted(pars, key=ref_to_note(notes))

    will sort the collection of (ref, paragraph) pairs according to the endnote text
    """

    def f(ref):
        return " ".join(
            map(lambda x: x.text, notes[ref[0].endnote_references[0].id].paragraphs)
        )

    return f


def copy_run_style(from_run, to_run, ignore=[], keep=[]):
    """
    Copies the style data from one run to another

    If you pass a list of styles to keep it ignore, it will use those lists to decide what to copy over i.e.
    it will copy keep - ignore.
    """

    copy_styles = [
        "all_caps",
        "bold",
        "complex_script",
        "cs_bold",
        "cs_italic",
        "double_strike",
        "emboss",
        "hidden",
        "italic",
        "imprint",
        "math",
        "no_proof",
        "outline",
        "rtl",
        "shadow",
        "small_caps",
        "snap_to_grid",
        "spec_vanish",
        "strike",
    ]
    if keep:
        copy_styles = filter(lambda x: x in keep, copy_styles)
    for style in filter(lambda x: x not in ignore, copy_styles):
        setattr(to_run, style, getattr(from_run, style))


def main(argv):
    parser = argparse.ArgumentParser()
    parser.add_argument("document", help="Document to load")
    parser.add_argument(
        "-n", type=int, default=10, help="Number of previous words to include"
    )
    parser.add_argument("-o", "--output", default="", help="Output file")
    parser.add_argument("--chapter", type=int, help="Load only refs from chapter")
    parser.add_argument("--sort", action="store_true")

    args = parser.parse_args(argv)

    if len(args.output) == 0:
        output_filename = args.document.rstrip(".docx") + "-endnotes.docx"
    else:
        output_filename = args.output

    d = Document(args.document)

    notes, note_ids = get_notes(d)
    pars = get_paragraphs(d, args.chapter)

    ids = set([])
    for run, _ in pars:
        ids.update(map(lambda x: x.id, run.endnote_references))

    # This is definitely doing something; there are only a few that don't match up
    assert ids.issubset(note_ids)

    d = Document()

    if args.sort:
        pars = sorted(pars, key=ref_to_note(notes))

    for ref, p in pars:
        id = ref.endnote_references[0].id
        runs = p.runs

        ix = ref_to_run_ix(id, runs)
        if ix == -1:
            print(f"Error: reference {id} not found in following paragraph:")
            print(p.text)
            sys.exit(1)

        pre_text = get_previous_words(ix, runs, args.n)

        for p2 in notes[id].paragraphs:
            np = d.add_paragraph()

            if len(pre_text) > 0 and len(p2.text.strip()) > 0:
                nr = np.add_run()
                nr.add_text("“" + pre_text + "”")
                nr.bold = True
                pre_text = ""  # Reset it so we only add it once

            for run in p2.runs:
                nr = np.add_run()
                nr.add_text(run.text)
                copy_run_style(run, nr, keep=["italic"])

        note_text = "".join(map(lambda x: x.text, notes[id].paragraphs))

        print(f"Note id: {id}")
        print(f"Note text: {note_text}")
        print(f"Source paragraph: {pre_text}")
        print("")

    d.save(output_filename)


if __name__ == "__main__":
    main(sys.argv[1:])
